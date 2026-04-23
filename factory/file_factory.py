"""具体解析器实现：支持 pdf, docx, doc, csv, xls, xlsx, txt, md"""
import os
import io
import subprocess
# import typing
import tempfile
from typing import Optional, List, Dict
from concurrent.futures import ThreadPoolExecutor, as_completed

# PDF - 使用 PyMuPDF (fitz)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

# DOCX
try:
    import docx
    from docx import Document
except Exception:
    docx = None

# DOC (binary) - 使用命令行工具解析，无需 textract

# CSV / Excel
try:
    import pandas as pd
except Exception:
    pd = None


def _parse_txt(data: bytes, encoding: Optional[str] = 'utf-8') -> str:
    try:
        return data.decode(encoding)
    except Exception:
        try:
            return data.decode('latin-1')
        except Exception:
            return data.decode('utf-8', errors='ignore')


def _parse_md(data: bytes) -> str:
    return _parse_txt(data)


def _parse_pdf(data: bytes) -> str:
    """
    使用 PyMuPDF 解析 PDF 文件并提取所有文本
    参数:
        data: PDF 文件的二进制数据
    返回:
        提取的完整文本内容
    异常:
        RuntimeError: 当 PyMuPDF 不可用时抛出
    """
    if not fitz:
        raise RuntimeError('PyMuPDF not installed or unavailable')
    try:
        # 从字节数据打开 PDF 文档
        doc = fitz.open(stream=data, filetype="pdf")
        texts = []
        # 遍历所有页面提取文本
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                texts.append(page_text)
        doc.close()
        return "\n\n".join(texts)
    except Exception as e:
        raise RuntimeError(f'PDF 解析失败：{str(e)}')


def _parse_pdf_with_pages(data: bytes) -> List[Dict]:
    """
    使用 PyMuPDF 解析 PDF 文件并按物理页面分割
    参数:
        data: PDF 文件的二进制数据
    返回:
        字典列表，每个字典包含:
        - content: 页面文本内容
        - page_number: 页码 (从 1 开始的整数)
    说明:
        PyMuPDF 按真实物理页面提取，避免了 pdfminer 的过度分割问题
        每页只提取有实际内容的文本，自动过滤空白页面
    """
    if not fitz:
        # PyMuPDF 不可用时回退到简单模式
        try:
            text = _parse_pdf(data)
            return [{"content": text, "page_number": None}]
        except Exception:
            return [{"content": "", "page_number": None}]
    try:
        # 从字节数据打开 PDF 文档
        doc = fitz.open(stream=data, filetype="pdf")
        pages = []
        # 遍历所有页面
        for page_num in range(len(doc)):
            page = doc[page_num]
            # 提取页面文本
            page_text = page.get_text("text").strip()
            # 只添加有实际内容的页面
            if page_text:
                pages.append({
                    "content": page_text,
                    "page_number": page_num + 1  # 页码从 1 开始
                })
        doc.close()
        # 如果所有页面都为空，返回空内容
        return pages if pages else [{"content": "", "page_number": None}]
    except Exception as e:
        print(f"PDF 分页解析失败：{e}，回退到单文本模式")
        try:
            text = _parse_pdf(data)
            return [{"content": text, "page_number": None}]
        except Exception:
            return [{"content": "", "page_number": None}]


def _parse_docx(data: bytes) -> str:
    """
    使用 python-docx 解析 DOCX 文件并提取所有文本
    参数:
        data: DOCX 文件的二进制数据
    返回:
        提取的完整文本内容 (已过滤控制字符)
    异常:
        RuntimeError: 当 python-docx 不可用时抛出
    """
    if docx is None:
        raise RuntimeError('python-docx not installed or unavailable')
    try:
        with io.BytesIO(data) as bio:
            doc = Document(bio)
            texts = []
            # 遍历所有段落提取文本
            for p in doc.paragraphs:
                try:
                    # 获取段落文本
                    para_text = p.text
                    # 确保是字符串类型
                    if para_text:
                        # 过滤控制字符 (保留换行符和制表符)
                        cleaned_text = ''.join(
                            char for char in para_text 
                            if ord(char) >= 32 or char in '\n\r\t'
                        )
                        # 只添加非空文本
                        if cleaned_text.strip():
                            texts.append(cleaned_text.strip())
                except Exception as e:
                    # 跳过有问题的段落
                    print(f"警告：跳过段落解析错误 - {e}")
                    continue
            # 合并所有文本
            return "\n\n".join(texts) if texts else ""
    except Exception as e:
        raise RuntimeError(f'DOCX 解析失败：{str(e)}')


def _parse_docx_with_pages(data: bytes) -> List[Dict]:
    """
    解析 DOCX 文件 (DOCX 本身没有页码概念，返回 None)
    参数:
        data: DOCX 文件的二进制数据
    返回:
        字典列表，每个字典包含:
        - content: 文本内容 (已过滤控制字符)
        - page_number: None (DOCX 无页码概念)
    说明:
        DOCX 文件没有物理页面概念，返回整个文档内容作为单个"页面"
        已添加控制字符过滤，避免乱码问题
    """
    if docx is None:
        raise RuntimeError('python-docx not installed or unavailable')
    try:
        with io.BytesIO(data) as bio:
            doc = Document(bio)
            texts = []
            # 遍历所有段落提取文本
            for p in doc.paragraphs:
                try:
                    para_text = p.text
                    if para_text:
                        # 过滤控制字符 (保留换行符和制表符)
                        cleaned_text = ''.join(
                            char for char in para_text 
                            if ord(char) >= 32 or char in '\n\r\t'
                        )
                        if cleaned_text.strip():
                            texts.append(cleaned_text.strip())
                except Exception as e:
                    print(f"警告：跳过段落解析错误 - {e}")
                    continue
            # 合并所有文本
            content = "\n\n".join(texts) if texts else ""
            # DOCX 文件没有页码概念，返回 None
            return [{"content": content, "page_number": None}]
    except Exception as e:
        print(f"DOCX 分页解析失败：{e}，回退到简单模式")
        try:
            content = _parse_docx(data)
            return [{"content": content, "page_number": None}]
        except Exception:
            return [{"content": "", "page_number": None}]


def _parse_doc(data: bytes) -> str:
    # .doc (binary) 使用命令行工具解析
    # 1. 尝试使用antiword命令行工具
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            # 尝试运行antiword
            result = subprocess.run(
                ['antiword', tmp_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            if result.returncode == 0:
                return result.stdout
        except (subprocess.SubprocessError, FileNotFoundError):
            # antiword不可用，继续尝试catdoc
            pass
        finally:
            # 清理临时文件
            try:
                os.unlink(tmp_path)
            except:
                pass
    except Exception:
        # 临时文件创建失败，继续尝试其他方法
        pass
    # 2. 尝试使用catdoc命令行工具
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            # 尝试运行catdoc
            result = subprocess.run(
                ['catdoc', tmp_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            if result.returncode == 0:
                return result.stdout
        except (subprocess.SubprocessError, FileNotFoundError):
            # catdoc也不可用
            pass
        finally:
            # 清理临时文件
            try:
                os.unlink(tmp_path)
            except:
                pass
    except Exception:
        # 临时文件创建失败
        pass
    # 3. 所有方法都失败，提供详细的错误信息
    error_msg = """
无法解析.doc文件。请安装以下任一命令行工具：
1. 安装antiword命令行工具：
   - Ubuntu/Debian: sudo apt-get install antiword
   - macOS: brew install antiword
   - Windows: 下载 http://www.winfield.demon.nl/ 并添加到PATH
   
2. 安装catdoc命令行工具：
   - Ubuntu/Debian: sudo apt-get install catdoc
   - macOS: brew install catdoc
   - Windows: 下载 http://www.winfield.demon.nl/ 并添加到PATH
   
3. 或者将.doc文件转换为.docx格式后再上传
"""
    raise RuntimeError(error_msg)


def _parse_doc_with_pages(data: bytes) -> List[Dict]:
    """
    解析 DOC 文件（DOC 本身没有可靠的页码概念，返回 None）
    返回：包含 content 和 page_number 的字典列表
    """
    content = _parse_doc(data)
    return [{"content": content, "page_number": None}]


def _parse_csv(data: bytes) -> str:
    if pd is None:
        raise RuntimeError('pandas not installed')
    with io.BytesIO(data) as bio:
        bio.seek(0)
        try:
            df = pd.read_csv(bio, dtype=str, engine='python')
        except Exception:
            bio.seek(0)
            df = pd.read_csv(bio, dtype=str, encoding='utf-8', engine='python', error_bad_lines=False)
        texts = []
        for r in df.fillna('').astype(str).values:
            texts.append(' '.join(r.tolist()))
        return '\n'.join(texts)


def _parse_csv_with_pages(data: bytes) -> List[Dict]:
    """
    解析 CSV 文件（无页码概念）
    返回：包含 content 和 page_number 的字典列表
    """
    content = _parse_csv(data)
    return [{"content": content, "page_number": None}]


def _parse_excel(data: bytes) -> str:
    if pd is None:
        raise RuntimeError('pandas not installed')
    with io.BytesIO(data) as bio:
        bio.seek(0)
        df = pd.read_excel(bio, sheet_name=None)
        texts = []
        for sheet, frame in df.items():
            texts.append(f"Sheet: {sheet}")
            for r in frame.fillna('').astype(str).values:
                texts.append(' '.join(r.tolist()))
        return '\n'.join(texts)


def _parse_excel_with_pages(data: bytes) -> List[Dict]:
    """
    解析 Excel 文件，每个 sheet 作为一个"页面"
    返回：包含 content 和 page_number 的字典列表
    """
    if pd is None:
        raise RuntimeError('pandas not installed')
    with io.BytesIO(data) as bio:
        bio.seek(0)
        df = pd.read_excel(bio, sheet_name=None)
        pages = []
        for i, (sheet, frame) in enumerate(df.items()):
            texts = []
            for r in frame.fillna('').astype(str).values:
                texts.append(' '.join(r.tolist()))
            content = f"Sheet: {sheet}\n" + '\n'.join(texts)
            pages.append({
                "content": content,
                "page_number": i + 1  # 使用 sheet 索引作为页码
            })
        return pages if pages else [{"content": "", "page_number": None}]


def _parse_txt_with_pages(data: bytes, encoding: Optional[str] = 'utf-8') -> List[Dict]:
    """
    解析 TXT 文件（TXT 文件没有页码概念，返回 None）
    返回：包含 content 和 page_number 的字典列表
    """
    try:
        text = data.decode(encoding)
    except Exception:
        try:
            text = data.decode('latin-1')
        except Exception:
            text = data.decode('utf-8', errors='ignore')
    # TXT 文件没有页码概念，直接返回整个内容，页码设为 None
    return [{"content": text, "page_number": None}]


def _parse_md_with_pages(data: bytes) -> List[Dict]:
    """
    解析 MD 文件（MD 文件没有页码概念，返回 None）
    返回：包含 content 和 page_number 的字典列表
    """
    try:
        text = data.decode('utf-8')
    except Exception:
        try:
            text = data.decode('latin-1')
        except Exception:
            text = data.decode('utf-8', errors='ignore')
    # MD 文件没有页码概念，直接返回整个内容，页码设为 None
    return [{"content": text, "page_number": None}]

PARSER_BY_EXT = {
    '.txt': _parse_txt,
    '.md': _parse_md,
    '.pdf': _parse_pdf,
    '.docx': _parse_docx,
    '.doc': _parse_doc,
    '.csv': _parse_csv,
    '.xls': _parse_excel,
    '.xlsx': _parse_excel,
}

# 带页码解析的映射表
PAGES_PARSER_BY_EXT = {
    '.txt': _parse_txt_with_pages,
    '.md': _parse_md_with_pages,
    '.pdf': _parse_pdf_with_pages,
    '.docx': _parse_docx_with_pages,
    '.doc': _parse_doc_with_pages,
    '.csv': _parse_csv_with_pages,
    '.xls': _parse_excel_with_pages,
    '.xlsx': _parse_excel_with_pages,
}


def get_parser_for_filename(filename: str):
    _, ext = os.path.splitext(filename.lower())
    return PARSER_BY_EXT.get(ext)


def get_pages_parser_for_filename(filename: str):
    """获取支持页码的解析器"""
    _, ext = os.path.splitext(filename.lower())
    return PAGES_PARSER_BY_EXT.get(ext)


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    parser = get_parser_for_filename(filename)
    if not parser:
        # fallback: try to decode as text
        return _parse_txt(data)
    return parser(data)


def extract_pages_from_bytes(filename: str, data: bytes) -> List[Dict]:
    """
    从文件字节数据中解析出带页码的页面内容
    参数:
        filename: 文件名（用于确定文件类型）
        data: 文件的二进制数据
    返回:
        字典列表，每个字典包含:
        - content: 页面内容
        - page_number: 页码（整数或 None）
    """
    parser = get_pages_parser_for_filename(filename)
    if not parser:
        # fallback: try to decode as text
        text = _parse_txt(data)
        return [{"content": text, "page_number": None}]
    try:
        return parser(data)
    except Exception as e:
        print(f"文件解析失败：{filename}, 错误：{e}")
        # 尝试回退到简单文本解析
        try:
            text = _parse_txt(data)
            return [{"content": text, "page_number": None}]
        except Exception:
            return [{"content": "", "page_number": None}]


def batch_extract_text_from_bytes(
    file_data_list: List[Dict[str, bytes]],
    max_workers: int = 4
) -> List[Dict]:
    """
    并发批量解析多个文件的文本内容
    参数:
        file_data_list: 文件数据列表，每个元素是 {'filename': str, 'data': bytes} 的字典
        max_workers: 最大并发线程数，默认4
    返回:
        解析结果列表，每个元素是 {
            'filename': str,
            'status': 'success' | 'failed',
            'content': str (成功时),
            'error': str (失败时)
        } 的字典
    """
    results = []
    
    def _parse_single_file(file_info: Dict[str, bytes]) -> Dict:
        """解析单个文件的内部函数"""
        filename = file_info['filename']
        data = file_info['data']
        try:
            content = extract_text_from_bytes(filename, data)
            return {
                'filename': filename,
                'status': 'success',
                'content': content
            }
        except Exception as e:
            return {
                'filename': filename,
                'status': 'failed',
                'error': str(e)
            }
    
    # 使用线程池并发解析
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {
            executor.submit(_parse_single_file, file_info): file_info
            for file_info in file_data_list
        }
        for future in as_completed(future_to_file):
            try:
                result = future.result()
                results.append(result)
            except Exception as e:
                file_info = future_to_file[future]
                results.append({
                    'filename': file_info['filename'],
                    'status': 'failed',
                    'error': f'线程执行异常: {str(e)}'
                })
    # 按原始顺序排序（可选）
    filename_order = {info['filename']: i for i, info in enumerate(file_data_list)}
    results.sort(key=lambda x: filename_order.get(x['filename'], 999))
    return results


def batch_extract_pages_from_bytes(
    file_data_list: List[Dict[str, bytes]],
    max_workers: int = 4
) -> List[Dict]:
    """
    并发批量解析多个文件的分页内容
    参数:
        file_data_list: 文件数据列表，每个元素是 {'filename': str, 'data': bytes} 的字典
        max_workers: 最大并发线程数，默认4
    返回:
        解析结果列表，每个元素是 {
            'filename': str,
            'status': 'success' | 'failed',
            'pages': List[Dict] (成功时),
            'error': str (失败时)
        } 的字典
    """
    results = []
    
    def _parse_single_file_pages(file_info: Dict[str, bytes]) -> Dict:
        """解析单个文件分页的内部函数"""
        filename = file_info['filename']
        data = file_info['data']
        try:
            pages = extract_pages_from_bytes(filename, data)
            return {
                'filename': filename,
                'status': 'success',
                'pages': pages
            }
        except Exception as e:
            return {
                'filename': filename,
                'status': 'failed',
                'error': str(e)
            }
    
    # 使用线程池并发解析
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {
            executor.submit(_parse_single_file_pages, file_info): file_info
            for file_info in file_data_list
        }
        for future in as_completed(future_to_file):
            try:
                result = future.result()
                results.append(result)
            except Exception as e:
                file_info = future_to_file[future]
                results.append({
                    'filename': file_info['filename'],
                    'status': 'failed',
                    'error': f'线程执行异常: {str(e)}'
                })
    # 按原始顺序排序
    filename_order = {info['filename']: i for i, info in enumerate(file_data_list)}
    results.sort(key=lambda x: filename_order.get(x['filename'], 999))
    return results
